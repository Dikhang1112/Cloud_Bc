import csv
from docx import Document
import re
from datetime import datetime


# Hàm chuẩn hóa header thành camelCase
def to_camel_case(text):
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)  # Loại bỏ ký tự đặc biệt
    words = text.lower().split()
    if not words:
        return ''
    return words[0] + ''.join(word.capitalize() for word in words[1:])


# Hàm kiểm tra tính hợp lệ của date
def is_valid_date(date_str):
    try:
        datetime.strptime(date_str.replace('/', '-'), '%d-%m-%Y')
        return True
    except ValueError:
        return False


# Hàm chính để xử lý ETL
def process_etl(input_file, output_file):
    try:
        # Đọc file Word
        doc = Document(input_file)
        if not doc.tables:
            raise Exception("File Word không chứa bảng dữ liệu")

        table = doc.tables[0]

        # Lấy header và chuẩn hóa
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        headers = [to_camel_case(header) for header in headers]

        # Đọc và transform dữ liệu
        data_rows = []
        error_log = []
        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_data = [cell.text.strip() for cell in row.cells]
            transformed_row = []

            for idx, cell_data in enumerate(row_data):
                if idx == 0:  # Cột nameProduct
                    # Loại bỏ ký tự đặc biệt, viết hoa chữ cái đầu
                    cleaned = re.sub(r'[^a-zA-Z\s]', '', cell_data).strip()
                    transformed_row.append(cleaned.title())
                elif idx in [1, 2]:  # Cột price và quantity
                    # Loại bỏ ký tự đặc biệt, chỉ giữ số
                    cleaned = re.sub(r'[^0-9-]', '', cell_data).strip()
                    # Chuyển thành số và lấy giá trị tuyệt đối (chuyển số âm thành dương)
                    try:
                        number = int(cleaned) if cleaned else 0
                        transformed_row.append(str(abs(number)))  # Chuyển số âm thành dương
                    except ValueError:
                        transformed_row.append(cleaned)  # Nếu không phải số, giữ nguyên
                elif idx == 3:  # Cột timeOfSale
                    cleaned = cell_data.replace('/', '-')
                    if is_valid_date(cleaned):
                        transformed_row.append(cleaned)
                    else:
                        error_log.append(f"Dòng {row_idx + 1}: Ngày không hợp lệ - {cleaned}")
                        transformed_row.append('INVALID_DATE')
                else:
                    transformed_row.append(cell_data)  # Các cột khác giữ nguyên

            data_rows.append(transformed_row)

        # Ghi dữ liệu đã transform vào file CSV
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(data_rows)

        # Ghi log lỗi (nếu có)
        if error_log:
            log_file = output_file.replace('.csv', '_error_log.txt')
            with open(log_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(error_log))
            return output_file, log_file
        return output_file, None

    except Exception as e:
        raise Exception(f"Không thể transform dữ liệu: {str(e)}")